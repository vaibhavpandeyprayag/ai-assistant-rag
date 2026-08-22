"""Format-specific document loading and text extraction.

Supported formats: PDF, DOCX, TXT, Markdown. Each loader converts a file into
a :class:`LoadedDocument` whose sections carry page numbers when the format
provides them. Third-party parser failures are converted to typed domain
errors so callers never see raw library exceptions.
"""

from pathlib import Path

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError as DocxPackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.errors import DocumentParseError, EmptyDocumentError, UnsupportedFormatError
from app.ingestion.models import DocumentSection, LoadedDocument

#: File extensions this module knows how to load.
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx", ".txt", ".md"})


def _load_pdf(path: Path) -> LoadedDocument:
    """Extract one section per PDF page, preserving 1-based page numbers."""
    try:
        reader = PdfReader(str(path))
        sections: list[DocumentSection] = []
        for page_index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            sections.append(DocumentSection(text=text.strip(), page_number=page_index))
    except PdfReadError as exc:
        raise DocumentParseError(
            f"Failed to parse PDF '{path.name}': the file appears to be corrupted."
        ) from exc

    return LoadedDocument(filename=path.name, source=str(path), sections=sections)


def _load_docx(path: Path) -> LoadedDocument:
    """Extract DOCX paragraph text; DOCX has no reliable page concept."""
    try:
        docx_file = DocxDocument(str(path))
    except DocxPackageNotFoundError as exc:
        raise DocumentParseError(
            f"Failed to parse DOCX '{path.name}': the file is not a valid Word document."
        ) from exc

    text = "\n".join(paragraph.text for paragraph in docx_file.paragraphs).strip()
    sections = [DocumentSection(text=text, page_number=None)]
    return LoadedDocument(filename=path.name, source=str(path), sections=sections)


def _load_plaintext(path: Path) -> LoadedDocument:
    """Read UTF-8 text for TXT and Markdown files."""
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentParseError(
            f"Failed to read '{path.name}': the file is not valid UTF-8 text."
        ) from exc

    sections = [DocumentSection(text=text.strip(), page_number=None)]
    return LoadedDocument(filename=path.name, source=str(path), sections=sections)


_LOADERS: dict[str, object] = {
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".txt": _load_plaintext,
    ".md": _load_plaintext,
}


def load_document(path: str | Path) -> LoadedDocument:
    """Load a document by dispatching on its file extension.

    Raises:
        FileNotFoundError: The path does not point at an existing file.
        UnsupportedFormatError: The extension is not supported.
        DocumentParseError: The file exists but cannot be parsed.
        EmptyDocumentError: The file parses but contains no extractable text.
    """
    document_path = Path(path)
    if not document_path.is_file():
        raise FileNotFoundError(f"File not found: {document_path}")

    suffix = document_path.suffix.lower()
    loader = _LOADERS.get(suffix)
    if loader is None:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFormatError(
            f"Unsupported file format '{suffix}'. Supported formats: {supported}."
        )

    loaded_document = loader(document_path)
    if not loaded_document.full_text.strip():
        raise EmptyDocumentError(
            f"Document '{document_path.name}' contains no extractable text."
        )
    return loaded_document
